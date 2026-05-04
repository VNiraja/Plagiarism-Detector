import os
import pdfplumber
from docx import Document

def extract_text_from_file(file_path):
    """
    Extracts text from PDF, DOCX, DOC, or TXT files.
    Returns the extracted text.
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return extract_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return extract_from_docx(file_path)
    elif file_ext == '.txt':
        return extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

def extract_from_pdf(file_path):
    """
    Extracts text from a PDF file.
    """
    text = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")
    
    return "\n".join(text)

def extract_from_docx(file_path):
    """
    Extracts text from a DOCX or DOC file.
    """
    text = []
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")
    
    return "\n".join(text)

def extract_from_txt(file_path):
    """
    Extracts text from a TXT file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Error reading TXT: {str(e)}")
